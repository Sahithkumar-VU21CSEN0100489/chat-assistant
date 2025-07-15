import os
from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Body, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi import Depends, status
from sqlalchemy.orm import Session
import models
import schemas
import database
import crud
import auth
import gemini
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
import shutil
from jose import jwt, JWTError
import PyPDF2
import docx
from schemas import AnswerResponse
import numpy as np
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import requests
from bs4 import BeautifulSoup
try:
    from readability import Document as ReadabilityDocument
except ImportError:
    ReadabilityDocument = None
import random
import smtplib
from email.mime.text import MIMEText
from pydantic import BaseModel

app = FastAPI()

# Allow CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

model = SentenceTransformer('all-MiniLM-L6-v2')

# In-memory OTP store (for demo only)
otp_store = {}

# Helper to send email (for demo, use your SMTP credentials)
def send_email(to_email, subject, body):
    smtp_server = 'smtp.gmail.com'
    smtp_port = 587
    smtp_user = 'sahith0489@gmail.com'
    smtp_pass = 'dzxk ehxk hyrm miom'
    msg = MIMEText(body)
    msg['Subject'] = subject
    msg['From'] = smtp_user
    msg['To'] = to_email
    with smtplib.SMTP(smtp_server, smtp_port) as server:
        server.starttls()
        server.login(smtp_user, smtp_pass)
        server.sendmail(smtp_user, [to_email], msg.as_string())

# Pydantic models for OTP endpoints
class EmailRequest(BaseModel):
    email: str
class OtpVerifyRequest(BaseModel):
    email: str
    otp: str
class ResetPasswordRequest(BaseModel):
    email: str
    otp: str
    new_password: str

@app.get("/")
def root():
    return {"message": "Backend is running"}

@app.get("/test-gemini")
async def test_gemini():
    try:
        result = await gemini.ask_gemini("The capital of France is Paris.", "What is the capital of France?")
        return {"result": result, "status": "success"}
    except Exception as e:
        return {"result": str(e), "status": "error"}

# Dependency

def get_db():
    db = database.SessionLocal()
    try:
        yield db
    finally:
        db.close()

database.init_db()

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/login")

# Refactored Auth endpoints
@app.post("/api/auth/signup", response_model=schemas.UserOut)
def signup(user: schemas.UserCreate, db: Session = Depends(get_db)):
    if crud.get_user_by_username(db, user.username):
        raise HTTPException(status_code=400, detail="Username already registered")
    user_obj = crud.create_user(db, user.username, user.email, user.password)
    return user_obj

@app.post("/api/auth/login")
def login(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    user = auth.authenticate_user(db, form_data.username, form_data.password)
    if not user:
        raise HTTPException(status_code=401, detail="Incorrect username or password")
    access_token = auth.create_access_token(data={"sub": user.username})
    return {"access_token": access_token, "token_type": "bearer", "user_id": user.id, "email": user.email}

@app.post("/api/auth/request-otp")
def request_otp(req: EmailRequest, db: Session = Depends(get_db)):
    email = req.email
    user = db.query(models.User).filter(models.User.email == email).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    otp = str(random.randint(100000, 999999))
    otp_store[email] = otp
    subject = "OTP for Login with Document Chat Assistant"
    body = f"Hello,\n\nYou requested to login with OTP for your account: {email}.\nYour OTP is: {otp}\n\nThis OTP is for login authentication. Do not share it with anyone.\n\nThanks,\nDocument Chat Assistant Team"
    send_email(email, subject, body)
    return {"message": "OTP sent to your email"}

@app.post("/api/auth/verify-otp")
def verify_otp(req: OtpVerifyRequest, db: Session = Depends(get_db)):
    email = req.email
    otp = req.otp
    if otp_store.get(email) != otp:
        raise HTTPException(status_code=400, detail="Invalid OTP")
    user = db.query(models.User).filter(models.User.email == email).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    # No email sent here, just verification
    access_token = auth.create_access_token(data={"sub": user.username})
    del otp_store[email]
    return {"access_token": access_token, "user_id": user.id, "email": user.email}

@app.post("/api/auth/reset-password")
def reset_password(req: ResetPasswordRequest, db: Session = Depends(get_db)):
    email = req.email
    otp = req.otp
    new_password = req.new_password
    if otp_store.get(email) != otp:
        raise HTTPException(status_code=400, detail="Invalid OTP")
    user = db.query(models.User).filter(models.User.email == email).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    user.hashed_password = auth.get_password_hash(new_password)
    db.add(user)
    db.commit()
    del otp_store[email]
    subject = "OTP for Password Reset - Document Chat Assistant"
    body = f"Hello,\n\nYou requested to reset your password for your account: {email}.\nYour OTP was: {otp}\n\nIf you did not request this, please contact support.\n\nThanks,\nDocument Chat Assistant Team"
    # Optionally, you can send a confirmation email here if needed
    return {"message": "Password reset successful"}

# Document upload
@app.post("/api/documents/upload", response_model=schemas.ChatOut)
def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    token: str = Depends(oauth2_scheme)
):
    print("[DEBUG] upload_document called")
    try:
        payload = jwt.decode(token, auth.SECRET_KEY, algorithms=[auth.ALGORITHM])
        print("[DEBUG] JWT decoded")
        username = payload.get("sub")
        print(f"[DEBUG] Username: {username}")
        if not username:
            print("[DEBUG] No username in token")
            raise HTTPException(status_code=401, detail="Invalid user")
        user = crud.get_user_by_username(db, username)
        print(f"[DEBUG] User: {user}")
        if not user:
            print("[DEBUG] User not found in DB")
            raise HTTPException(status_code=401, detail="Invalid user")
        # Restrict to 3 documents per user
        user_id = user.id  # user.id is already the correct type for DB queries
        user_documents = crud.get_documents_by_user(db, user_id)
        if len(user_documents) >= 3:
            raise HTTPException(
                status_code=400,
                detail="You can only upload up to 3 documents. Delete existing documents to upload new ones."
            )
        content = None
        print(f"[DEBUG] File content type: {file.content_type}")
        if file.content_type == "text/plain":
            print("Processing text file")
            try:
                content = file.file.read().decode("utf-8")
            except UnicodeDecodeError:
                file.file.seek(0)  # Reset file pointer
                try:
                    content = file.file.read().decode("latin-1")
                except UnicodeDecodeError:
                    file.file.seek(0)
                    content = file.file.read().decode("cp1252", errors="ignore")
            print("Finished text file")
        elif file.content_type == "application/pdf":
            print("Processing PDF file")
            reader = PyPDF2.PdfReader(file.file)
            content = "\n".join(page.extract_text() or "" for page in reader.pages)
            # --- Table Extraction for PDF ---
            tables = []
            for page in reader.pages:
                text = page.extract_text() or ""
                # Simple heuristic: look for lines with multiple columns separated by spaces/tabs
                lines = text.split("\n")
                for line in lines:
                    if "\t" in line or (line.count(" ") > 5):
                        tables.append(line)
            if tables:
                content += "\n\n[Extracted Tables from PDF]\n" + "\n".join(tables)
            print("Finished PDF file with table extraction")
        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            print("Processing DOCX file")
            docx_file = docx.Document(file.file)
            content = "\n".join([para.text for para in docx_file.paragraphs])
            # --- Table Extraction for DOCX ---
            docx_tables = []
            for table in docx_file.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    docx_tables.append(" | ".join(row_text))
            if docx_tables:
                content += "\n\n[Extracted Tables from DOCX]\n" + "\n".join(docx_tables)
            print("Finished DOCX file with table extraction")
        else:
            print("Unsupported file type:", file.content_type)
            raise HTTPException(status_code=400, detail="Unsupported file type.")
        if not content or not content.strip():
            print("[DEBUG] Uploaded file is empty")
            raise HTTPException(status_code=400, detail="Uploaded file is empty.")
        filename = file.filename or "uploaded.txt"
        print(f"[DEBUG] Creating document: {filename}")
        user_id = user.id  # user.id is already the correct type for DB queries
        if not isinstance(user_id, int):
            raise HTTPException(status_code=500, detail="User ID is invalid")
        doc = crud.create_document(db, filename, content, user_id)
        print(f"[DEBUG] Document created with id: {doc.id}")
        # --- RAG: Chunk and embed document ---
        chunk_size = 500
        chunks = [content[i:i+chunk_size] for i in range(0, len(content), chunk_size)]
        print(f"[DEBUG] Number of chunks: {len(chunks)}")
        # Fix: Add error handling, type checks, and ensure embeddings and chunks match
        try:
            embeddings = model.encode(chunks)
            print("[DEBUG] Embeddings generated")
        except Exception as e:
            print(f"[DEBUG] Embedding generation failed: {e}")
            raise HTTPException(status_code=500, detail="Failed to generate embeddings.")

        doc_id = getattr(doc, "id", None)
        if doc_id is None or not isinstance(doc_id, int):
            print("[DEBUG] Invalid document ID")
            raise HTTPException(status_code=500, detail="Document ID is invalid")

        if not isinstance(embeddings, (list, np.ndarray)) or len(embeddings) != len(chunks):
            print("[DEBUG] Embeddings and chunks length mismatch")
            raise HTTPException(status_code=500, detail="Embedding and chunk count mismatch.")

        for i, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
            try:
                chunk = crud.create_document_chunk(
                    db,
                    doc_id,
                    chunk_text,
                    np.array(embedding).tobytes()
                )
            except Exception as e:
                print(f"[DEBUG] Failed to create chunk {i+1}: {e}")
                raise HTTPException(status_code=500, detail=f"Failed to create document chunk {i+1}.")
            print(f"[DEBUG] Created chunk {i+1}/{len(chunks)} with id: {chunk.id}")
        print("[DEBUG] All chunks stored")
        chat = crud.create_chat(db, doc_id)
        print(f"[DEBUG] Chat created with id: {chat.id}")
        return chat
    except JWTError:
        print("[DEBUG] JWTError: Invalid or expired token")
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    except Exception as e:
        print(f"[DEBUG] Exception: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/links/upload", response_model=schemas.ChatOut)
async def upload_link(
    request: Request,
    db: Session = Depends(get_db),
    token: str = Depends(oauth2_scheme)
):
    data = await request.json()
    url = data.get("url")
    if not url:
        raise HTTPException(status_code=400, detail="URL is required.")
    try:
        payload = jwt.decode(token, auth.SECRET_KEY, algorithms=[auth.ALGORITHM])
        username = payload.get("sub")
        if not username:
            raise HTTPException(status_code=401, detail="Invalid user")
        user = crud.get_user_by_username(db, username)
        if not user:
            raise HTTPException(status_code=401, detail="Invalid user")
        # Restrict to 3 documents per user
        user_id = getattr(user, "id", None)
        if user_id is None:
            raise HTTPException(status_code=500, detail="User ID is missing")
        try:
            user_id = int(user_id)
        except (TypeError, ValueError):
            raise HTTPException(status_code=500, detail="User ID is invalid")
        user_documents = crud.get_documents_by_user(db, user_id)
        if len(user_documents) >= 3:
            raise HTTPException(status_code=400, detail="You can only upload up to 3 documents. Delete existing documents to upload new ones.")
        # Fetch the URL content
        try:
            resp = requests.get(url, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
            resp.raise_for_status()
            html = resp.text
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to fetch URL: {e}")
        # Extract main content using readability-lxml and BeautifulSoup
        try:
            if ReadabilityDocument is None:
                raise HTTPException(status_code=500, detail="readability-lxml is not installed on the server.")
            doc = ReadabilityDocument(html)
            main_html = doc.summary() if doc else ""
            if not main_html:
                raise HTTPException(status_code=400, detail="Could not extract main content from the webpage.")
            soup_main = BeautifulSoup(main_html, "html.parser")
            main_content = soup_main.get_text(separator="\n", strip=True)
            soup_full = BeautifulSoup(html, "html.parser")
            full_content = soup_full.get_text(separator="\n", strip=True)
            # Merge main and full content, remove duplicates
            combined = main_content + "\n" + full_content
            lines = [line.strip() for line in combined.split("\n") if line.strip()]
            # Remove duplicates while preserving order
            seen = set()
            content = "\n".join([x for x in lines if not (x in seen or seen.add(x))])
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to extract main content: {e}")
        # Clean and validate content
        content = content.strip()
        if not content or len(content) < 100:
            raise HTTPException(status_code=400, detail="Extracted content is too short or empty. The website may not be supported.")
        if len(content) > 20000:
            content = content[:20000]  # Limit to 20,000 characters
            filename = url
            # Ensure user.id is a plain int, not a SQLAlchemy column/expression
            user_id = getattr(user, "id", None)
            if user_id is None:
                raise HTTPException(status_code=500, detail="User ID is missing")
            try:
                user_id = int(user_id)
            except (TypeError, ValueError):
                raise HTTPException(status_code=500, detail="User ID is invalid")
            # Store as a document with fileType 'link'
            doc = crud.create_document(db, filename, content, user_id)
            # --- RAG: Chunk and embed document ---
            chunk_size = 500
        chunks = [content[i:i+chunk_size] for i in range(0, len(content), chunk_size)]
        embeddings = model.encode(chunks)
        # Ensure doc.id is a plain int, not a SQLAlchemy column/expression
        doc_id = getattr(doc, "id", None)
        if doc_id is None:
            raise HTTPException(status_code=500, detail="Document ID is missing")
        try:
            doc_id = int(doc_id)
        except (TypeError, ValueError):
            raise HTTPException(status_code=500, detail="Document ID is invalid")
        for i, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
            chunk = crud.create_document_chunk(db, doc_id, chunk_text, np.array(embedding).tobytes())
        chat = crud.create_chat(db, doc_id)
        return chat
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# List all documents for the authenticated user
@app.get("/api/documents", response_model=list[schemas.DocumentOut])
def list_documents(db: Session = Depends(get_db), token: str = Depends(oauth2_scheme)):
    payload = jwt.decode(token, auth.SECRET_KEY, algorithms=[auth.ALGORITHM])
    username = payload.get("sub")
    if not username:
        raise HTTPException(status_code=401, detail="Invalid user")
    user = crud.get_user_by_username(db, username)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid user")
    # Ensure user.id is a plain int, not a SQLAlchemy column/expression
    user_id = getattr(user, "id", None)
    if user_id is None:
        raise HTTPException(status_code=500, detail="User ID is missing")
    try:
        user_id = int(user_id)
    except (TypeError, ValueError):
        raise HTTPException(status_code=500, detail="User ID is invalid")
    documents = crud.get_documents_by_user(db, user_id)
    return documents or []
@app.get("/api/documents/{document_id}", response_model=schemas.DocumentOut)
def get_document(document_id: int, db: Session = Depends(get_db), token: str = Depends(oauth2_scheme)):
    try:
        payload = jwt.decode(token, auth.SECRET_KEY, algorithms=[auth.ALGORITHM])
        username = payload.get("sub")
        if not username:
            raise HTTPException(status_code=401, detail="Invalid user")
        user = crud.get_user_by_username(db, username)
        if not user:
            raise HTTPException(status_code=401, detail="Invalid user")
        doc = crud.get_document_by_id(db, document_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        if getattr(doc, "owner_id", None) != getattr(user, "id", None):
            raise HTTPException(status_code=403, detail="Not authorized to access this document")
        return doc
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid or expired token")

# Delete a document
@app.delete("/api/documents/{document_id}")
def delete_document(document_id: int, db: Session = Depends(get_db), token: str = Depends(oauth2_scheme)):
    try:
        payload = jwt.decode(token, auth.SECRET_KEY, algorithms=[auth.ALGORITHM])
        username = payload.get("sub")
        if not username:
            raise HTTPException(status_code=401, detail="Invalid user")
        user = crud.get_user_by_username(db, username)
        if not user:
            raise HTTPException(status_code=401, detail="Invalid user")
        doc = crud.get_document_by_id(db, document_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        if getattr(doc, "owner_id", None) != getattr(user, "id", None):
            raise HTTPException(status_code=403, detail="Not authorized to delete this document")
        crud.delete_document(db, document_id)
        return {"detail": "Document deleted"}
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid or expired token")

# Chat/QA endpoint
@app.post("/api/chat/{document_id}/ask", response_model=schemas.AnswerResponse)
async def ask_question(document_id: int, req: schemas.QuestionRequest, db: Session = Depends(get_db), token: str = Depends(oauth2_scheme)):
    try:
        payload = jwt.decode(token, auth.SECRET_KEY, algorithms=[auth.ALGORITHM])
        username = payload.get("sub")
        if not username:
            raise HTTPException(status_code=401, detail="Invalid user")
        user = crud.get_user_by_username(db, username)
        if not user:
            raise HTTPException(status_code=401, detail="Invalid user")
        doc = crud.get_document_by_id(db, document_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        if getattr(doc, "owner_id", None) != getattr(user, "id", None):
            raise HTTPException(status_code=403, detail="Not authorized to access this document")
        gemini_result = await gemini.ask_gemini(getattr(doc, "content", ""), req.question)
        return {"answer": gemini_result["answer"], "followup_questions": gemini_result["followup_questions"]}
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid or expired token")

@app.post("/api/ask", response_model=AnswerResponse)
async def ask_document_question(
    question: str = Body(...),
    document_id: int = Body(...),
    db: Session = Depends(get_db),
    token: str = Depends(oauth2_scheme)
):
    try:
        payload = jwt.decode(token, auth.SECRET_KEY, algorithms=[auth.ALGORITHM])
        username = payload.get("sub")
        if not username:
            raise HTTPException(status_code=401, detail="Invalid user")
        user = crud.get_user_by_username(db, username)
        if not user:
            raise HTTPException(status_code=401, detail="Invalid user")
        user_id = getattr(user, "id", None)
        if user_id is None:
            raise HTTPException(status_code=500, detail="User ID is invalid")
        print(f"[DEBUG] Looking for document {document_id} owned by user {user_id}")
        document = db.query(models.Document).filter(
            models.Document.id == document_id,
            models.Document.owner_id == user_id
        ).first()
        if not document:
            print(f"[DEBUG] Document {document_id} not found or not owned by user {user_id}")
            raise HTTPException(status_code=404, detail="Document not found")
        # --- RAG retrieval ---
        doc_id = document.id  # document.id should already be an int
        if not isinstance(doc_id, int):
            raise HTTPException(status_code=500, detail="Document ID is invalid")
        chunks = crud.get_chunks_by_document(db, doc_id)
        debug_info = f"Retrieved {len(chunks)} chunks for document {doc_id}"
        print(f"[DEBUG] {debug_info}")
        if not chunks:
            debug_info += "\nNo chunks found in database"
            print("[DEBUG] No chunks found in database")
            return {"answer": f"No information available. Debug: {debug_info}", "followup_questions": []}
        q_lower = question.lower()
        is_summary = ("summary" in q_lower) or ("summar" in q_lower)
        if is_summary:
            context_parts = [getattr(chunk, "chunk_text", "") for chunk in chunks[:20]]
            context = "\n".join(context_parts)
            debug_info += f"\n[SUMMARY MODE] Sent {len(context_parts)} chunks for summary."
        else:
            question_emb = model.encode([question])[0]
            chunk_embs = []
            for chunk in chunks:
                embedding = chunk.embedding
                if hasattr(embedding, "value"):
                    embedding = embedding.value
                if isinstance(embedding, (bytes, bytearray)):
                    try:
                        chunk_emb = np.frombuffer(embedding, dtype=np.float32)
                    except Exception:
                        chunk_emb = np.frombuffer(bytes(embedding), dtype=np.float32)
                    chunk_embs.append(chunk_emb)
            sims = cosine_similarity([question_emb], chunk_embs)[0]
            debug_info += f"\nSimilarity scores: {sims}"
            debug_info += f"\nMax similarity score: {np.max(sims)}"
            debug_info += f"\nMin similarity score: {np.min(sims)}"
            print(f"[DEBUG] Similarity scores: {sims}")
            print(f"[DEBUG] Max similarity score: {np.max(sims)}")
            print(f"[DEBUG] Min similarity score: {np.min(sims)}")
            top_indices = np.argsort(sims)[::-1][:15]  # Top 15 chunks
            debug_info += f"\nTop chunk indices: {top_indices}"
            debug_info += f"\nTop chunk scores: {[sims[i] for i in top_indices]}"
            print(f"[DEBUG] Top chunk indices: {top_indices}")
            print(f"[DEBUG] Top chunk scores: {[sims[i] for i in top_indices]}")
            top_chunks = [chunks[i] for i in top_indices if sims[i] > 0.01]
            debug_info += f"\nChunks above threshold (0.01): {len(top_chunks)}"
            print(f"[DEBUG] Chunks above threshold (0.01): {len(top_chunks)}")
            if not top_chunks:
                debug_info += "\nNo chunks above threshold. Using top 10 chunks regardless of score."
                print("[DEBUG] No chunks above threshold. Using top 10 chunks regardless of score.")
                top_chunks = [chunks[i] for i in top_indices if i < len(chunks)][:10]
                if not top_chunks:
                    debug_info += "\nStill no chunks available."
                    print("[DEBUG] Still no chunks available. Returning no information.")
                    return {"answer": f"No information available. Debug: {debug_info}", "followup_questions": []}
            context_parts = []
            for chunk in top_chunks:
                chunk_text = getattr(chunk, "chunk_text", None)
                if chunk_text is None and isinstance(chunk, (list, tuple)) and len(chunk) > 0:
                    if chunk_text is None and isinstance(chunk, (list, tuple)) and len(chunk) > 0:
                        chunk_text = getattr(chunk[0], "chunk_text", None)
                    if chunk_text is not None and hasattr(chunk_text, "value"):
                        chunk_text = chunk_text.value
                    context_parts.append(str(chunk_text) if chunk_text is not None else "")
                else:
                    context_parts.append(str(chunk_text) if chunk_text is not None else "")
            context = "\n".join(context_parts)
            debug_info += f"\nContext length: {len(context)}"
            debug_info += f"\nContext preview: {context[:200]}..."
            print(f"[DEBUG] Context sent to Gemini (first 500 chars): {context[:500]}")
        gemini_result = await gemini.ask_gemini(context, question)
        return {"answer": gemini_result["answer"], "followup_questions": gemini_result["followup_questions"]}
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    except Exception as e:
        import traceback
        print("[DEBUG] Exception in /api/ask:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))